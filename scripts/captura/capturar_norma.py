"""C.3.a — Captura responsável de uma URL (lei/decreto/programa).

Implementa @.claude/rules/captura-responsavel.md (R1–R11). Para cada URL:
  1. Verifica robots.txt (cache 24h)
  2. Aguarda rate-limit (1 req/2s; respeita Crawl-delay)
  3. GET com User-Agent identificável + timeouts (10/30/60)
  4. Validação bruta (200, tamanho mínimo, sem regex de erro)
  5. Hash SHA-256 dos bytes; dedupe content-addressable
  6. Salva snapshot em data/external_snapshots/<sha[:2]>/<sha>.<ext>
  7. Extrai texto: trafilatura (HTML), pdfplumber (PDF), python-docx (DOCX)
  8. Validação extraída (tamanho ≥ 100 chars; PII scan CPF/CNPJ flag)
  9. Metadata em data/extracted_text/<sha>.metadata.json
 10. Log JSONL em data/logs/captura_<YYYY-MM-DD>.jsonl
 11. Atualiza data/external_snapshots/index.json

Uso programático:
    from capturar_norma import capturar
    resultado = capturar("https://...", tipo_documento="lei")

Uso CLI:
    python -B scripts/captura/capturar_norma.py <url> [--tipo lei|decreto|...]
"""
from __future__ import annotations

import argparse
import hashlib
import io
import json
import logging
import os
import re
import subprocess
import sys
import time
from dataclasses import asdict, dataclass
from datetime import datetime, timezone, timedelta
from pathlib import Path
from urllib.parse import urlparse

import httpx

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent.parent
sys.path.insert(0, str(HERE))
from _http_helpers import RobotsCache, RateLimiter, make_client, USER_AGENT, timeout_for
from _external_tools import tesseract_bin, soffice_bin, tessdata_prefix, has_tesseract_pt

LOG = logging.getLogger(__name__)

DATA_HOJE = datetime.now().strftime("%Y-%m-%d")
TZ_BR = timezone(timedelta(hours=-3))

SNAPSHOTS_DIR = ROOT / "data" / "external_snapshots"
EXTRACTED_DIR = ROOT / "data" / "extracted_text"
LOG_PATH = ROOT / "data" / "logs" / f"captura_{DATA_HOJE}.jsonl"
INDEX_PATH = SNAPSHOTS_DIR / "index.json"

REGEX_PAGINA_ERRO = re.compile(
    r"(página\s+não\s+encontrada|erro\s+404|acesso\s+negado|page\s+not\s+found)",
    re.IGNORECASE,
)
REGEX_CPF = re.compile(r"\d{3}\.\d{3}\.\d{3}-\d{2}")
REGEX_CNPJ = re.compile(r"\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}")

ATRIBUICAO_POR_DOMINIO = {
    "planalto.gov.br": "Brasil. Presidência da República. Casa Civil.",
    "in.gov.br": "Diário Oficial da União — Imprensa Nacional",
    "camara.leg.br": "Câmara dos Deputados",
    "senado.leg.br": "Senado Federal",
    "mec.gov.br": "Ministério da Educação",
    "inep.gov.br": "INEP — Ministério da Educação",
    "gov.br": "Governo Federal — gov.br (CC BY-ND 3.0)",
}


@dataclass
class CapturaResultado:
    status: str  # ok | bloqueado_robots | falha_status | falha_rede | validacao_falhou | inalterado
    url_solicitada: str
    url_final: str | None = None
    http_status: int | None = None
    content_type: str | None = None
    sha256: str | None = None
    extensao: str | None = None
    tamanho_bytes: int | None = None
    tempo_ms: int = 0
    encoding: str | None = None
    caminho_snapshot: str | None = None
    caminho_texto: str | None = None
    caracteres_extraidos: int | None = None
    ocr_aplicado: bool = False
    contem_pii: bool = False
    pii_count: int = 0
    atribuicao: str | None = None
    licenca_inferida: str | None = None
    metodo_http: str | None = None  # GET | HEAD+GET (fallback)
    retries: int = 0
    erro_tipo: str | None = None
    erro_msg: str | None = None
    timestamp_iso: str | None = None


def infer_atribuicao(url: str) -> str:
    host = urlparse(url).netloc.lower()
    for sub, atr in ATRIBUICAO_POR_DOMINIO.items():
        if sub in host:
            return atr
    return f"Conteúdo público — {host} (atribuição inferida)"


def infer_licenca(tipo: str | None, url: str) -> str:
    if tipo in {"lei", "decreto", "portaria", "instrucao_normativa", "resolucao"}:
        return "dominio_publico_lei_8_iv"
    if "gov.br" in url.lower():
        return "CC BY-ND 3.0 (presumida; gov.br)"
    return "sem_licenca_explicita"


def infer_extensao(content_type: str | None, url: str) -> str:
    if content_type:
        ct = content_type.lower()
        if "html" in ct:
            return "html"
        if "pdf" in ct:
            return "pdf"
        if "msword" in ct or "officedocument.wordprocessingml" in ct:
            return "docx"
        if "opendocument" in ct:
            return "odt"
        if "json" in ct:
            return "json"
        if "xml" in ct:
            return "xml"
        if "plain" in ct:
            return "txt"
    # fallback por URL
    p = urlparse(url).path.lower()
    for ext in ("html", "htm", "pdf", "docx", "doc", "odt"):
        if p.endswith("." + ext):
            return ext if ext != "htm" else "html"
    return "html"  # default razoável


def extract_text(content: bytes, ext: str, encoding: str | None = None) -> str:
    """Extrai texto de bytes conforme extensão (sem OCR — usar extract_with_ocr_fallback)."""
    if ext in ("html", "htm"):
        try:
            import trafilatura
            txt = trafilatura.extract(
                content,
                favor_recall=True,
                include_links=False,
                include_tables=True,
                output_format="txt",
            )
            if txt:
                return txt
        except Exception:
            pass
        # Fallback: bs4
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, "lxml")
            for tag in soup(["script", "style"]):
                tag.decompose()
            return soup.get_text(separator="\n", strip=True)
        except Exception:
            return ""
    if ext == "pdf":
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                return "\n".join((p.extract_text() or "") for p in pdf.pages)
        except Exception as e:
            LOG.warning(f"pdfplumber falhou: {e}")
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(content))
                return "\n".join((p.extract_text() or "") for p in reader.pages)
            except Exception:
                return ""
    if ext == "docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs)
            # Adiciona texto de tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += "\n" + cell.text
            return text
        except Exception:
            return ""
    if ext in ("doc",):
        # DOC legado → libreoffice headless converte para txt
        return extract_doc_via_libreoffice(content) or ""
    if ext == "odt":
        try:
            from odf.opendocument import load
            from odf import text as odf_text, teletype
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".odt", delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            try:
                doc = load(tmp_path)
                paragraphs = doc.getElementsByType(odf_text.P)
                return "\n".join(teletype.extractText(p) for p in paragraphs)
            finally:
                Path(tmp_path).unlink(missing_ok=True)
        except Exception as e:
            LOG.warning(f"odfpy falhou: {e}")
            return ""
    if ext == "txt":
        try:
            return content.decode(encoding or "utf-8", errors="replace")
        except Exception:
            return ""
    return ""


def extract_pdf_with_ocr_fallback(
    content: bytes,
    sha: str,
    *,
    min_chars_para_ocr: int = 100,
) -> tuple[str, bool]:
    """Tenta pdfplumber primeiro; se texto < min_chars_para_ocr, aplica ocrmypdf + Tesseract pt.

    Retorna (texto, ocr_aplicado).
    """
    texto = extract_text(content, "pdf")
    if len(texto) >= min_chars_para_ocr:
        return texto, False

    # Fallback OCR
    if not has_tesseract_pt():
        LOG.warning(f"Tesseract pt não disponível; pulando OCR para sha={sha[:12]}")
        return texto, False

    try:
        import ocrmypdf
        import tempfile
        with tempfile.TemporaryDirectory() as td:
            in_pdf = Path(td) / f"{sha}.pdf"
            out_pdf = Path(td) / f"{sha}.ocr.pdf"
            in_pdf.write_bytes(content)

            # Configurar TESSDATA_PREFIX se houver pasta local
            prefix = tessdata_prefix()
            old_env = os.environ.get("TESSDATA_PREFIX")
            if prefix:
                os.environ["TESSDATA_PREFIX"] = str(prefix)
            try:
                ocrmypdf.ocr(
                    str(in_pdf), str(out_pdf),
                    language="por",
                    skip_text=False,
                    force_ocr=True,
                    deskew=True,
                    progress_bar=False,
                    output_type="pdf",
                )
            finally:
                if old_env is None:
                    os.environ.pop("TESSDATA_PREFIX", None)
                else:
                    os.environ["TESSDATA_PREFIX"] = old_env

            # Re-extrair com pdfplumber sobre o PDF com camada OCR
            ocr_bytes = out_pdf.read_bytes()
            texto_ocr = extract_text(ocr_bytes, "pdf")
            if len(texto_ocr) > len(texto):
                LOG.info(f"OCR aplicado em sha={sha[:12]}: {len(texto)} → {len(texto_ocr)} chars")
                return texto_ocr, True
    except Exception as e:
        LOG.warning(f"ocrmypdf falhou para sha={sha[:12]}: {e}")

    return texto, False


def extract_doc_via_libreoffice(content: bytes, timeout: int = 30) -> str | None:
    """Converte .doc legado para .txt via LibreOffice headless. Retorna texto ou None."""
    bin_ = soffice_bin()
    if not bin_:
        LOG.warning("LibreOffice não disponível; não é possível processar .doc")
        return None
    try:
        import tempfile
        with tempfile.TemporaryDirectory() as td:
            in_doc = Path(td) / "in.doc"
            in_doc.write_bytes(content)
            r = subprocess.run(
                [str(bin_), "--headless", "--convert-to", "txt", "--outdir", td, str(in_doc)],
                capture_output=True, text=True, encoding="utf-8",
                timeout=timeout,
            )
            if r.returncode != 0:
                LOG.warning(f"soffice falhou (rc={r.returncode}): {r.stderr[:200]}")
                return None
            out_txt = Path(td) / "in.txt"
            if not out_txt.is_file():
                return None
            return out_txt.read_text(encoding="utf-8", errors="replace")
    except subprocess.TimeoutExpired:
        LOG.warning("soffice timeout em conversão .doc")
        return None
    except Exception as e:
        LOG.warning(f"soffice exception: {e}")
        return None


def detect_encoding(content: bytes) -> str:
    try:
        from charset_normalizer import detect
        return (detect(content)["encoding"] or "utf-8")
    except Exception:
        return "utf-8"


def update_snapshot_index(resultado: "CapturaResultado") -> None:
    """Atualiza data/external_snapshots/index.json com info do snapshot capturado.

    Estrutura:
      {
        "by_sha": { "<sha>": { url_original, url_canonica, extensao, data_captura, ... } },
        "by_url": { "<url_canonica>": "<sha>" },
        "ultima_atualizacao": "..."
      }
    """
    if not resultado.sha256:
        return  # nada a indexar
    INDEX_PATH.parent.mkdir(parents=True, exist_ok=True)
    if INDEX_PATH.exists():
        try:
            idx = json.loads(INDEX_PATH.read_text(encoding="utf-8"))
        except Exception:
            idx = {"by_sha": {}, "by_url": {}, "ultima_atualizacao": None}
    else:
        idx = {"by_sha": {}, "by_url": {}, "ultima_atualizacao": None}

    sha = resultado.sha256
    url_canon = resultado.url_final or resultado.url_solicitada

    entry = idx["by_sha"].get(sha, {})
    entry.update({
        "url_original": resultado.url_solicitada,
        "url_canonica": url_canon,
        "extensao": resultado.extensao,
        "data_captura": entry.get("data_captura") or resultado.timestamp_iso,
        "ultima_validacao": resultado.timestamp_iso,
        "tamanho_bytes": resultado.tamanho_bytes,
        "ocr_aplicado": resultado.ocr_aplicado,
        "atribuicao": resultado.atribuicao,
        "licenca_inferida": resultado.licenca_inferida,
        "caracteres_extraidos": resultado.caracteres_extraidos,
        "contem_pii": resultado.contem_pii,
    })
    idx["by_sha"][sha] = entry
    idx["by_url"][url_canon] = sha
    idx["ultima_atualizacao"] = resultado.timestamp_iso

    INDEX_PATH.write_text(json.dumps(idx, ensure_ascii=False, indent=2), encoding="utf-8")


def capturar(
    url: str,
    *,
    tipo_documento: str | None = None,
    client: httpx.Client | None = None,
    robots_cache: RobotsCache | None = None,
    rate_limiter: RateLimiter | None = None,
) -> CapturaResultado:
    """Captura uma URL conforme R1–R11. Idempotente (dedupe por SHA-256)."""
    timestamp = datetime.now(TZ_BR).isoformat(timespec="seconds")
    resultado = CapturaResultado(
        status="indefinido",
        url_solicitada=url,
        timestamp_iso=timestamp,
    )

    own_client = client is None
    own_robots = robots_cache is None
    own_rate = rate_limiter is None
    if own_client:
        client = make_client()
    if own_robots:
        robots_cache = RobotsCache()
    if own_rate:
        rate_limiter = RateLimiter()

    try:
        # R2: robots.txt
        host = urlparse(url).netloc.lower()
        robots = robots_cache.check(url, client)
        if robots.crawl_delay:
            rate_limiter.set_delay(host, robots.crawl_delay)
        if not robots.allowed:
            resultado.status = "bloqueado_robots"
            return resultado

        # R3: rate limit
        rate_limiter.wait(url)

        # R4: GET com timeout específico do host (planalto.gov.br tem 90s read)
        host_timeout = timeout_for(url)
        t0 = time.monotonic()
        retries = 0

        def _fetch(method: str = "GET", attempt_timeout: httpx.Timeout = host_timeout) -> httpx.Response:
            return client.request(method, url, follow_redirects=True, timeout=attempt_timeout)

        try:
            r = _fetch("GET")
            resultado.metodo_http = "GET"
        except httpx.TimeoutException as e:
            # Retry para hosts lentos: 1 tentativa adicional após backoff curto
            if "planalto.gov.br" in url.lower():
                rate_limiter.wait(url)
                time.sleep(15)
                retries += 1
                try:
                    r = _fetch("GET")
                    resultado.metodo_http = "GET (retry)"
                except httpx.HTTPError as e2:
                    resultado.status = "falha_rede"
                    resultado.erro_tipo = type(e2).__name__
                    resultado.erro_msg = str(e2)[:200]
                    resultado.tempo_ms = int((time.monotonic() - t0) * 1000)
                    resultado.retries = retries
                    return resultado
            else:
                resultado.status = "falha_rede"
                resultado.erro_tipo = "TimeoutException"
                resultado.erro_msg = str(e)[:200]
                resultado.tempo_ms = int((time.monotonic() - t0) * 1000)
                return resultado
        except httpx.HTTPError as e:
            resultado.status = "falha_rede"
            resultado.erro_tipo = type(e).__name__
            resultado.erro_msg = str(e)[:200]
            resultado.tempo_ms = int((time.monotonic() - t0) * 1000)
            return resultado

        # 403 → tentar User-Agent alternativo? Não — postura ética é aceitar.
        # Mas se o erro for HTTP 405/501 (method not allowed em GET — improvável),
        # tentar HEAD apenas (não vai capturar conteúdo, só validar)
        if r.status_code in (405, 501):
            try:
                r2 = _fetch("HEAD")
                if r2.status_code in (200, 204):
                    # Servidor aceita HEAD mas não GET — peculiar; aceitar status
                    r = r2
                    resultado.metodo_http = "HEAD (fallback de 405/501)"
                    retries += 1
            except httpx.HTTPError:
                pass

        resultado.tempo_ms = int((time.monotonic() - t0) * 1000)
        resultado.url_final = str(r.url)
        resultado.http_status = r.status_code
        resultado.content_type = r.headers.get("content-type", "")
        resultado.retries = retries

        # R7: validação bruta
        if r.status_code != 200:
            resultado.status = "falha_status"
            resultado.erro_msg = f"HTTP {r.status_code}"
            return resultado

        content = r.content
        resultado.tamanho_bytes = len(content)

        ext = infer_extensao(resultado.content_type, url)
        # Tamanho mínimo
        min_bytes = 1024 if ext in ("html", "htm") else 5120
        if resultado.tamanho_bytes < min_bytes:
            resultado.status = "validacao_falhou"
            resultado.erro_msg = f"Tamanho {resultado.tamanho_bytes}B abaixo do mínimo {min_bytes}B"
            return resultado

        # R5: snapshot content-addressable
        sha = hashlib.sha256(content).hexdigest()
        resultado.sha256 = sha
        resultado.extensao = ext
        snap_subdir = SNAPSHOTS_DIR / sha[:2]
        snap_subdir.mkdir(parents=True, exist_ok=True)
        snap_path = snap_subdir / f"{sha}.{ext}"

        if snap_path.exists():
            # Dedupe: já temos. Apenas atualizar metadata se necessário.
            resultado.status = "inalterado"
        else:
            snap_path.write_bytes(content)
            resultado.status = "ok"
        try:
            resultado.caminho_snapshot = str(snap_path.relative_to(ROOT)).replace("\\", "/")
        except ValueError:
            # Snapshot fora da pasta do projeto (ex.: em testes com tmp_path) — usar absoluto
            resultado.caminho_snapshot = str(snap_path).replace("\\", "/")

        # Encoding
        encoding = detect_encoding(content)
        resultado.encoding = encoding

        # Extração de texto (com OCR fallback para PDF)
        if ext == "pdf":
            texto, ocr_aplicado = extract_pdf_with_ocr_fallback(content, sha)
            resultado.ocr_aplicado = ocr_aplicado
        else:
            texto = extract_text(content, ext, encoding)
        resultado.caracteres_extraidos = len(texto) if texto else 0

        # Validação extraída: regex erro
        if texto and REGEX_PAGINA_ERRO.search(texto[:5000]):
            resultado.status = "validacao_falhou"
            resultado.erro_msg = "Texto contém marcador de página de erro"
            # ainda salvamos snapshot e texto para auditoria

        # PII scan
        cpfs = REGEX_CPF.findall(texto or "")
        cnpjs = REGEX_CNPJ.findall(texto or "")
        resultado.pii_count = len(cpfs) + len(cnpjs)
        resultado.contem_pii = resultado.pii_count > 5  # threshold de R8

        # Salvar texto extraído + metadata
        EXTRACTED_DIR.mkdir(parents=True, exist_ok=True)
        text_path = EXTRACTED_DIR / f"{sha}.txt"
        if texto and not resultado.contem_pii:
            text_path.write_text(texto, encoding="utf-8")
            try:
                resultado.caminho_texto = str(text_path.relative_to(ROOT)).replace("\\", "/")
            except ValueError:
                resultado.caminho_texto = str(text_path).replace("\\", "/")

        # R9: atribuição
        resultado.atribuicao = infer_atribuicao(url)
        resultado.licenca_inferida = infer_licenca(tipo_documento, url)

        meta_path = EXTRACTED_DIR / f"{sha}.metadata.json"
        meta = {
            **asdict(resultado),
            "tipo_documento_declarado": tipo_documento,
            "extensao": ext,
        }
        meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

        # Atualiza index.json (idempotente; merge com entry existente)
        update_snapshot_index(resultado)

        return resultado

    finally:
        # R10: log JSONL sempre
        try:
            LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
            with LOG_PATH.open("a", encoding="utf-8") as fh:
                fh.write(json.dumps(asdict(resultado), ensure_ascii=False) + "\n")
        except Exception:
            pass

        if own_client:
            client.close()


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("url")
    parser.add_argument("--tipo", default=None, help="lei | decreto | portaria | ...")
    args = parser.parse_args()

    res = capturar(args.url, tipo_documento=args.tipo)
    print(json.dumps(asdict(res), ensure_ascii=False, indent=2))
    return 0 if res.status in ("ok", "inalterado") else 2


if __name__ == "__main__":
    sys.exit(main())
